#!/usr/bin/env python3
"""
Proposal Checker — validates AI-generated proposal documents against a local library.

Usage:
    python proposal_checker.py <proposal.docx> --library <folder> [--rules proposal_rules.yaml] [--output checked.docx]

Highlights:
    Green  = paragraph verified against library source
    Yellow = contains unverified factual claim (number, cert, %)
    Red    = rule violation (forbidden phrase, missing required content)
"""

import argparse
import os
import re
import sys
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional

# --- Optional dependency checks with helpful messages ---

try:
    from docx import Document
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)

try:
    import yaml
except ImportError:
    yaml = None

try:
    import pandas as pd
except ImportError:
    pd = None

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from rapidfuzz import fuzz
except ImportError:
    fuzz = None


# ---------------------------------------------------------------------------
# Data structures
# ---------------------------------------------------------------------------

@dataclass
class LibraryEntry:
    source_file: str
    location: str  # e.g. "row 5", "page 2", "paragraph 12"
    text: str


@dataclass
class CheckResult:
    para_index: int   # -1 for document-level checks
    para_text: str
    status: str       # "pass" | "warn" | "fail"
    issues: list = field(default_factory=list)
    sources: list = field(default_factory=list)


# ---------------------------------------------------------------------------
# Library indexing — reads all supported file types from a folder
# ---------------------------------------------------------------------------

def index_library(library_path: str) -> list:
    lib = Path(library_path)
    if not lib.exists():
        print(f"ERROR: Library folder not found: {library_path}")
        sys.exit(1)

    entries = []
    for filepath in lib.rglob("*"):
        if not filepath.is_file():
            continue
        ext = filepath.suffix.lower()
        if ext == ".docx":
            entries.extend(_index_docx(filepath))
        elif ext in (".xlsx", ".xls", ".csv"):
            entries.extend(_index_spreadsheet(filepath))
        elif ext == ".pdf":
            entries.extend(_index_pdf(filepath))
        elif ext in (".txt", ".md"):
            entries.extend(_index_text(filepath))

    return entries


def _index_docx(filepath: Path) -> list:
    try:
        doc = Document(str(filepath))
        return [
            LibraryEntry(filepath.name, f"paragraph {i+1}", p.text.strip())
            for i, p in enumerate(doc.paragraphs)
            if p.text.strip()
        ]
    except Exception as e:
        print(f"  Warning: could not read {filepath.name}: {e}")
        return []


def _index_spreadsheet(filepath: Path) -> list:
    if pd is None:
        print(f"  Warning: pandas not installed — skipping {filepath.name}")
        return []
    try:
        entries = []
        if filepath.suffix.lower() == ".csv":
            df = pd.read_csv(filepath, dtype=str)
            for i, row in df.iterrows():
                text = " | ".join(str(v) for v in row.values if pd.notna(v) and str(v).strip())
                if text:
                    entries.append(LibraryEntry(filepath.name, f"row {i+2}", text))
        else:
            xl = pd.ExcelFile(filepath)
            for sheet in xl.sheet_names:
                df = xl.parse(sheet, dtype=str)
                for i, row in df.iterrows():
                    text = " | ".join(str(v) for v in row.values if pd.notna(v) and str(v).strip())
                    if text:
                        entries.append(LibraryEntry(filepath.name, f"{sheet} row {i+2}", text))
        return entries
    except Exception as e:
        print(f"  Warning: could not read {filepath.name}: {e}")
        return []


def _index_pdf(filepath: Path) -> list:
    if pdfplumber is None:
        print(f"  Warning: pdfplumber not installed — skipping {filepath.name}")
        return []
    try:
        entries = []
        with pdfplumber.open(str(filepath)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text() or ""
                for line in text.splitlines():
                    if line.strip():
                        entries.append(LibraryEntry(filepath.name, f"page {i+1}", line.strip()))
        return entries
    except Exception as e:
        print(f"  Warning: could not read {filepath.name}: {e}")
        return []


def _index_text(filepath: Path) -> list:
    try:
        text = filepath.read_text(encoding="utf-8", errors="ignore")
        return [
            LibraryEntry(filepath.name, f"line {i+1}", line.strip())
            for i, line in enumerate(text.splitlines())
            if line.strip()
        ]
    except Exception as e:
        print(f"  Warning: could not read {filepath.name}: {e}")
        return []


# ---------------------------------------------------------------------------
# Claim extraction — finds verifiable facts in a paragraph
# ---------------------------------------------------------------------------

CLAIM_PATTERNS = {
    "certification":    r"\b(?:ISO\s*\d+(?::\d+)?|CMMI\s*(?:Level\s*)?\d|DCAA[\s-]?(?:approved|compliant)?|FedRAMP|SOC\s*[12]|ITAR|FISMA|DoD\s*\d+\.\d+|NIST\s*\d+(?:\.\d+)*|CAGE\s*(?:code\s*)?\w{5})\b",
    "percentage":       r"\b\d+(?:\.\d+)?%",
    "dollar_amount":    r"\$[\d,]+(?:\.\d+)?(?:\s*(?:million|billion|thousand|[MBK]))?\b",
    "years_experience": r"\b\d+\+?\s+years?\b",
    "employee_count":   r"\b\d[\d,]*\s+(?:employees?|staff|personnel|FTEs?)\b",
    "contract_number":  r"\b[A-Z]{1,4}\d{3,}-\d{2}-[A-Z]-\d{4,}\b",
    "naics_code":       r"\bNAICS\s*(?:code\s*)?\d{6}\b",
    "duns_uei":         r"\b(?:DUNS|UEI)[:\s]*[A-Z0-9]{9,13}\b",
}


def extract_claims(text: str) -> list:
    """Return [(claim_type, matched_text), ...] for verifiable claims in text."""
    claims = []
    for claim_type, pattern in CLAIM_PATTERNS.items():
        for match in re.finditer(pattern, text, re.IGNORECASE):
            claims.append((claim_type, match.group()))
    return claims


# ---------------------------------------------------------------------------
# Library search — finds entries that support a claim
# ---------------------------------------------------------------------------

def search_library(claim: str, library: list, threshold: int = 80) -> list:
    """Return [(LibraryEntry, score), ...] for entries that match the claim."""
    claim_lower = claim.lower()
    results = []

    if fuzz is None:
        # Substring fallback
        for entry in library:
            if claim_lower in entry.text.lower():
                results.append((entry, 100))
    else:
        for entry in library:
            score = fuzz.partial_ratio(claim_lower, entry.text.lower())
            if score >= threshold:
                results.append((entry, score))

    results.sort(key=lambda x: x[1], reverse=True)
    return results[:3]


# ---------------------------------------------------------------------------
# Rules loading
# ---------------------------------------------------------------------------

DEFAULT_RULES = {
    "required_sections":          [],
    "required_phrases":           [],
    "forbidden_phrases":          [],
    "claim_verification_enabled": True,
    "fuzzy_threshold":            80,
}


def load_rules(rules_file: Optional[str]) -> dict:
    if not rules_file or not Path(rules_file).exists():
        if rules_file:
            print(f"  Rules file not found ({rules_file}), using defaults.")
        return DEFAULT_RULES.copy()

    if yaml is None:
        print("  Warning: PyYAML not installed — rules file ignored. Run: pip install pyyaml")
        return DEFAULT_RULES.copy()

    with open(rules_file, encoding="utf-8") as f:
        loaded = yaml.safe_load(f) or {}

    rules = DEFAULT_RULES.copy()
    rules.update(loaded)
    return rules


# ---------------------------------------------------------------------------
# Proposal checking
# ---------------------------------------------------------------------------

def check_proposal(proposal_path: str, library: list, rules: dict) -> list:
    doc = Document(proposal_path)
    results = []

    found_sections = set()
    full_text = "\n".join(p.text for p in doc.paragraphs)

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        result = CheckResult(para_index=i, para_text=text, status="pass")

        # Track required section headings
        for section in rules.get("required_sections", []):
            if section.lower() in text.lower():
                found_sections.add(section)

        # Forbidden phrases
        for phrase in rules.get("forbidden_phrases", []):
            if phrase.lower() in text.lower():
                result.status = "fail"
                result.issues.append(f"Forbidden phrase: \"{phrase}\"")

        # Factual claim verification
        if rules.get("claim_verification_enabled", True):
            for claim_type, claim_text in extract_claims(text):
                matches = search_library(claim_text, library, rules.get("fuzzy_threshold", 80))
                if matches:
                    best_entry, best_score = matches[0]
                    result.sources.append(
                        f"{claim_text}  →  {best_entry.source_file} ({best_entry.location}, match {best_score}%)"
                    )
                else:
                    if result.status == "pass":
                        result.status = "warn"
                    result.issues.append(f"Unverified {claim_type}: \"{claim_text}\" — not found in library")

        results.append(result)

    # Document-level: missing required sections
    for section in rules.get("required_sections", []):
        if section not in found_sections:
            results.append(CheckResult(
                para_index=-1,
                para_text="[DOCUMENT-LEVEL]",
                status="fail",
                issues=[f"Required section missing: \"{section}\""],
            ))

    # Document-level: required phrases absent from full document
    for phrase in rules.get("required_phrases", []):
        if phrase.lower() not in full_text.lower():
            results.append(CheckResult(
                para_index=-1,
                para_text="[DOCUMENT-LEVEL]",
                status="fail",
                issues=[f"Required phrase missing from document: \"{phrase}\""],
            ))

    return results


# ---------------------------------------------------------------------------
# Annotated .docx output
# ---------------------------------------------------------------------------

# Highlight fill colors (hex, no #)
COLOR_PASS = "C6EFCE"   # light green
COLOR_WARN = "FFEB9C"   # light yellow
COLOR_FAIL = "FFC7CE"   # light red
COLOR_NOTE = "F2F2F2"   # light gray (annotation rows)


def _apply_shading(para_elem, fill_hex: str):
    pPr = para_elem.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    # Remove existing shd if present
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    pPr.append(shd)


def _make_annotation_element(text: str, fill_hex: str):
    """Build a <w:p> XML element for an annotation row (not yet attached)."""
    p = OxmlElement("w:p")

    pPr = OxmlElement("w:pPr")
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)
    p.append(pPr)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "404040")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "16")   # 8 pt
    sz_cs = OxmlElement("w:szCs")
    sz_cs.set(qn("w:val"), "16")
    i_el = OxmlElement("w:i")

    rPr.extend([color_el, sz, sz_cs, i_el])
    r.append(rPr)

    t = OxmlElement("w:t")
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    t.text = text
    r.append(t)
    p.append(r)

    return p


def write_annotated_doc(proposal_path: str, results: list, output_path: str):
    doc = Document(proposal_path)

    result_map = {r.para_index: r for r in results if r.para_index >= 0}
    doc_level = [r for r in results if r.para_index == -1]

    color_for_status = {"pass": COLOR_PASS, "warn": COLOR_WARN, "fail": COLOR_FAIL}

    # Iterate paragraphs in reverse so addnext insertions don't shift indices
    for i in range(len(doc.paragraphs) - 1, -1, -1):
        result = result_map.get(i)
        if result is None:
            continue

        para = doc.paragraphs[i]
        _apply_shading(para._p, color_for_status.get(result.status, "FFFFFF"))

        if result.issues or result.sources:
            label = {"pass": "PASS", "warn": "NEEDS REVIEW", "fail": "FAIL"}.get(result.status, "?")
            lines = [f"[{label}]"]
            for issue in result.issues:
                lines.append(f"  Issue: {issue}")
            for src in result.sources:
                lines.append(f"  Source: {src}")
            annotation_text = "  |  ".join(lines)

            ann_color = COLOR_NOTE if result.status == "pass" else color_for_status[result.status]
            ann_elem = _make_annotation_element(annotation_text, ann_color)
            para._p.addnext(ann_elem)

    # Append document-level issues as a summary table at the end
    if doc_level:
        doc.add_paragraph()
        hdr = doc.add_paragraph("DOCUMENT-LEVEL REVIEW RESULTS")
        hdr.runs[0].bold = True
        _apply_shading(hdr._p, COLOR_FAIL)

        for r in doc_level:
            for issue in r.issues:
                p = doc.add_paragraph(f"    {issue}")
                _apply_shading(p._p, COLOR_FAIL)

    doc.save(output_path)
    print(f"Annotated document saved: {output_path}")


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Check an AI-generated proposal .docx against a local source library.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python proposal_checker.py draft_proposal.docx --library C:/proposals/library
  python proposal_checker.py draft.docx --library ./library --rules proposal_rules.yaml --output draft_checked.docx
        """,
    )
    parser.add_argument("proposal", help="Path to the proposal .docx file to check")
    parser.add_argument("--library", required=True, help="Folder containing source library files (xlsx, docx, pdf, txt, csv)")
    parser.add_argument("--rules", default="proposal_rules.yaml", help="Rules YAML config file (default: proposal_rules.yaml)")
    parser.add_argument("--output", help="Output path for annotated .docx (default: <name>_checked.docx)")
    args = parser.parse_args()

    if not Path(args.proposal).exists():
        print(f"ERROR: Proposal file not found: {args.proposal}")
        sys.exit(1)

    output_path = args.output or (Path(args.proposal).stem + "_checked.docx")

    print(f"\n{'='*60}")
    print("  PROPOSAL CHECKER")
    print(f"{'='*60}")
    print(f"  Proposal : {args.proposal}")
    print(f"  Library  : {args.library}")
    print(f"  Rules    : {args.rules}")
    print(f"  Output   : {output_path}")
    print(f"{'='*60}\n")

    print("[1/3] Loading rules...")
    rules = load_rules(args.rules)
    print(f"      Required sections  : {len(rules.get('required_sections', []))}")
    print(f"      Required phrases   : {len(rules.get('required_phrases', []))}")
    print(f"      Forbidden phrases  : {len(rules.get('forbidden_phrases', []))}")

    print("\n[2/3] Indexing library...")
    library = index_library(args.library)
    print(f"      {len(library)} entries indexed")

    if not library:
        print("  WARNING: Library is empty. Claim verification will flag everything as unverified.")

    print("\n[3/3] Checking proposal...")
    results = check_proposal(args.proposal, library, rules)

    # Print summary
    passes = sum(1 for r in results if r.status == "pass")
    warns  = sum(1 for r in results if r.status == "warn")
    fails  = sum(1 for r in results if r.status == "fail")

    print(f"\n{'='*60}")
    print(f"  RESULTS:  {passes} passed  |  {warns} warnings  |  {fails} failures")
    print(f"{'='*60}")

    if warns or fails:
        print("\nIssues:")
        for r in results:
            if r.issues:
                loc = f"Para {r.para_index+1}" if r.para_index >= 0 else "Document"
                preview = r.para_text[:70] + ("..." if len(r.para_text) > 70 else "")
                print(f"\n  [{r.status.upper()}] {loc}: {preview}")
                for issue in r.issues:
                    print(f"    - {issue}")

    print()
    write_annotated_doc(args.proposal, results, output_path)


if __name__ == "__main__":
    main()
