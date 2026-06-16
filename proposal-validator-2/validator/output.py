"""
Output generation from Claude Code's results.json.

build_highlighted_docx(results)             -> bytes (.docx)
build_json_report(results, proposal, folder) -> dict (provenance report)
"""
import logging
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX

logger = logging.getLogger(__name__)

_HIGHLIGHT = {
    "GREEN": WD_COLOR_INDEX.GREEN,
    "RED": WD_COLOR_INDEX.RED,
    "YELLOW": WD_COLOR_INDEX.YELLOW,
}


def build_highlighted_docx(results: dict) -> bytes:
    """
    Reconstruct the full proposal as a .docx with per-verdict highlight colours.
    Non-factual paragraphs (verdict=null) get no highlight.
    """
    doc = Document()
    # Remove the default empty paragraph python-docx inserts
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    for para in results.get("paragraphs", []):
        text = para.get("text") or ""
        verdict = para.get("verdict")
        highlight = _HIGHLIGHT.get(verdict) if verdict else None

        # Preserve any embedded newlines as separate Word paragraphs
        for line in text.split("\n"):
            wp = doc.add_paragraph()
            run = wp.add_run(line)
            if highlight is not None:
                run.font.highlight_color = highlight

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_json_report(results: dict, proposal_path: str, source_folder: str) -> dict:
    """Reformat Claude Code's results into the provenance report schema."""
    summary = results.get("summary", {})
    factual = [p for p in results.get("paragraphs", []) if p.get("is_factual")]

    return {
        "proposal_file": Path(proposal_path).name,
        "source_folder": str(source_folder),
        "run_timestamp": results.get(
            "completed_at", datetime.now(timezone.utc).isoformat()
        ),
        "summary": {
            "total_claims": summary.get("total_claims", len(factual)),
            "green": summary.get("green", 0),
            "red": summary.get("red", 0),
            "yellow": summary.get("yellow", 0),
            "non_factual": summary.get("non_factual", 0),
        },
        "claims": [
            {
                "claim_id": p.get("span_id"),
                "claim_text": p.get("text"),
                "claim_type": p.get("claim_type"),
                "verdict": p.get("verdict"),
                "confidence": p.get("confidence"),
                "source_file": p.get("source_file"),
                "source_type": p.get("source_type"),
                "source_location": p.get("source_location"),
                "source_excerpt": p.get("source_excerpt"),
                "explanation": p.get("explanation"),
            }
            for p in factual
        ],
    }
